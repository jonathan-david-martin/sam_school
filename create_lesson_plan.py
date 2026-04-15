from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Header ---
title = doc.add_heading('The Sam School — Lesson Plan', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Grades 5–6  |  April 15, 2026  |  Science & Technology')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# --- Unit Context ---
doc.add_heading('Unit Context', level=1)
p = doc.add_paragraph('Students have just completed a unit covering:')
bullets = [
    'Size of planets (absolute measurements)',
    'Relative scale of planets (comparing sizes)',
    'Unit conversion (e.g., miles ↔ kilometers, feet ↔ meters)',
    'Gravity (how it differs across planets, effect on weight)'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()

# ============================================================
# PERIOD 1
# ============================================================
doc.add_heading('Period 1: Build a Space Science Video Game', level=1)

# Overview
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'Students will use Claude Code to design and build an educational web-based video game '
    'that reinforces the concepts from our planets and gravity unit. Each team will prompt '
    'Claude Code to generate a playable game, test it locally, and then we will deploy '
    'the best version for the school to play.'
)

# Learning Objectives
doc.add_heading('Learning Objectives', level=2)
objectives = [
    'Demonstrate understanding of planet sizes and relative scale',
    'Apply unit conversion skills within a game context',
    'Explain how gravity differs across planets',
    'Practice computational thinking and iterative design',
    'Use AI tools (Claude Code) as a creative collaborator'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# Materials
doc.add_heading('Materials', level=2)
materials = [
    'Laptops/Chromebooks with Claude Code installed (1 per team)',
    'Projector for demos',
    'Whiteboard for brainstorming',
    'Game Design Worksheet (see appendix)'
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

# Lesson Flow
doc.add_heading('Lesson Flow', level=2)

# Table for lesson flow
table = doc.add_table(rows=6, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Time', 'Activity', 'Details']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

rows_data = [
    ['5 min', 'Warm-Up & Hook',
     'Show a quick example of a simple browser game. Ask: "What if YOU could build a game '
     'that teaches other kids about the solar system?" Introduce today\'s challenge.'],
    ['10 min', 'Game Design Workshop',
     'Teams brainstorm on the Game Design Worksheet:\n'
     '• What type of game? (quiz, platformer, space explorer, sorting game)\n'
     '• Which concepts will it teach? (must include ALL: planet sizes, unit conversion, gravity)\n'
     '• What makes it fun? (scoring, levels, surprises)\n'
     'Each team presents their concept in 30 seconds.'],
    ['20 min', 'Build with Claude Code',
     'Teams open Claude Code and prompt it to build their game. Teacher circulates to help '
     'with prompting strategies.\n\n'
     'STARTER PROMPT for students:\n'
     '"Build me a web game using HTML, CSS, and JavaScript. The game should teach players '
     'about [their concept]. Include [specific features]. Make it colorful and fun for kids."\n\n'
     'Students iterate: playtest → give feedback to Claude → improve → repeat.'],
    ['8 min', 'Playtest & Peer Review',
     'Teams swap laptops and playtest each other\'s games. Testers fill out feedback:\n'
     '• Did you learn something about planets/gravity/units?\n'
     '• What was the most fun part?\n'
     '• One suggestion to make it better.'],
    ['2 min', 'Wrap-Up & Voting',
     'Quick class vote on which game(s) to deploy. Teacher will deploy selected games '
     'after class (or live if time permits).'],
]

for i, row_data in enumerate(rows_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph()

# Game Ideas
doc.add_heading('Sample Game Ideas for Students', level=2)
ideas = [
    'Planet Sorter — Drag and drop planets into the correct size order. Bonus round: convert their diameters from miles to kilometers.',
    'Gravity Jumper — A platformer where your character jumps on different planets. Jump height changes based on each planet\'s gravity. Display your weight on each planet.',
    'Space Quiz Show — Timed quiz covering all unit topics. Questions get harder each round. Leaderboard tracks scores.',
    'Unit Conversion Race — Players race to convert measurements before time runs out. Space-themed with planet facts between rounds.',
    'Build-a-Solar-System — Place planets at the correct relative distances and sizes. Earn points for accuracy.'
]
for idea in ideas:
    doc.add_paragraph(idea, style='List Bullet')

# Prompting Tips
doc.add_heading('Prompting Tips for Students', level=2)
tips = [
    'Be specific: "Make the planets the correct relative sizes" works better than "add planets."',
    'Iterate: "The quiz is good but add a timer and make wrong answers show the correct fact."',
    'Ask for fixes: "The button doesn\'t work when I click it. Can you fix it?"',
    'Add polish last: "Now make it look really cool with a space background and animations."'
]
for t in tips:
    doc.add_paragraph(t, style='List Bullet')

# Deployment
doc.add_heading('Deployment Plan', level=2)
doc.add_paragraph(
    'After the class selects the best game(s), the teacher will deploy using one of these options:'
)
deploy = [
    'GitHub Pages — Push the HTML/CSS/JS files to a GitHub repo and enable Pages (free, instant)',
    'Netlify Drop — Drag and drop the game folder onto netlify.com/drop for instant hosting',
    'School intranet — Copy files to the school\'s shared server if available'
]
for d in deploy:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph()

# ============================================================
# PERIOD 2
# ============================================================
doc.add_heading('Period 2: VEX Go Competition Prep', level=1)

# Overview
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'Students will prepare for the upcoming VEX Go robotics competition by designing strategies, '
    'building/modifying their robots, and programming them. Claude Code will be used to help '
    'students think through their programming logic and generate pseudocode that they implement '
    'in VEXcode GO\'s block-based environment.'
)

# Note about VEX Go
doc.add_heading('Important Note: VEX Go & Claude Code', level=2)
doc.add_paragraph(
    'VEX Go uses VEXcode GO, a block-based programming app (no text-based coding). '
    'Claude Code cannot directly generate VEXcode GO block programs, but it CAN:'
)
vex_uses = [
    'Generate pseudocode and step-by-step logic that students translate into blocks',
    'Help students think through strategy and problem-solving ("If the robot hits a wall, what should it do?")',
    'Explain programming concepts (loops, conditionals, sensor logic) in plain language',
    'Create decision flowcharts for robot behavior',
    'Help debug: students describe what\'s going wrong and Claude suggests fixes'
]
for v in vex_uses:
    doc.add_paragraph(v, style='List Bullet')

# Learning Objectives
doc.add_heading('Learning Objectives', level=2)
objectives2 = [
    'Develop a competition strategy based on game rules',
    'Program robot movements using loops, conditionals, and sensor input',
    'Use AI tools to plan and debug robot logic',
    'Practice teamwork and communication under time constraints',
    'Iterate on designs through testing and refinement'
]
for obj in objectives2:
    doc.add_paragraph(obj, style='List Bullet')

# Materials
doc.add_heading('Materials', level=2)
materials2 = [
    'VEX Go robot kits (1 per team)',
    'Tablets/Chromebooks with VEXcode GO app installed',
    'Laptops with Claude Code (1 per team for strategy/pseudocode)',
    'Competition field or practice area',
    'Competition rules printout',
    'Strategy Planning Worksheet (see appendix)'
]
for m in materials2:
    doc.add_paragraph(m, style='List Bullet')

# Lesson Flow
doc.add_heading('Lesson Flow', level=2)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Medium Shading 1 Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h

rows_data2 = [
    ['5 min', 'Rally & Rules Review',
     'Quick review of competition rules and scoring. Show a short clip of a VEX Go match. '
     'Ask: "What strategies do you notice the winning teams using?"'],
    ['5 min', 'Strategy Session with Claude Code',
     'Teams use Claude Code to brainstorm strategy:\n\n'
     'SAMPLE PROMPT:\n'
     '"We\'re in a VEX Go competition. Our robot has motors, a bumper sensor, an eye sensor, '
     'and an electromagnet. The challenge is [describe competition task]. '
     'Help us plan a step-by-step strategy and pseudocode for our robot."\n\n'
     'Claude Code outputs a logical plan that students will implement in blocks.'],
    ['20 min', 'Build, Code & Test',
     'Teams split responsibilities:\n'
     '• Builders: modify/improve robot hardware\n'
     '• Coders: translate Claude\'s pseudocode into VEXcode GO blocks\n'
     '• Strategists: set up practice runs, track what works\n\n'
     'Cycle: Code → Test on field → Ask Claude for debugging help → Improve → Repeat'],
    ['10 min', 'Practice Matches',
     'Run 2–3 timed practice matches. Teams rotate opponents. '
     'After each match, 2 minutes for teams to huddle and adjust.\n'
     'Encourage students to ask Claude Code: "Our robot keeps [problem]. What should we change?"'],
    ['5 min', 'Debrief & Goal Setting',
     'Each team shares:\n'
     '• One thing that worked well\n'
     '• One thing to improve before competition\n'
     '• Their #1 goal for next practice session\n'
     'Teacher records goals for follow-up.'],
]

for i, row_data in enumerate(rows_data2):
    for j, cell_text in enumerate(row_data):
        table2.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph()

# Sample Claude Code prompts for VEX Go
doc.add_heading('Sample Claude Code Prompts for VEX Go', level=2)
prompts = [
    '"Explain what a loop does in programming. Give me a robot example a 5th grader would understand."',
    '"My robot needs to drive forward until it hits something, then turn right and keep going. Write this as pseudocode."',
    '"The eye sensor detects red and blue objects. Write pseudocode for: if red, pick up with electromagnet; if blue, turn around."',
    '"Our robot keeps turning too far when we tell it to turn 90 degrees. What could be causing this and how do we fix it?"',
    '"Create a flowchart for our robot\'s competition strategy: start → drive to objects → sort by color → return to base."'
]
for pr in prompts:
    doc.add_paragraph(pr, style='List Bullet')

doc.add_paragraph()

# ============================================================
# APPENDIX
# ============================================================
doc.add_heading('Appendix A: Game Design Worksheet', level=1)

doc.add_paragraph('Team Name: _______________________')
doc.add_paragraph('Team Members: _______________________')
doc.add_paragraph()

questions = [
    '1. Game Title: _______________________',
    '2. Type of game (circle one):   Quiz   |   Platformer   |   Sorting/Matching   |   Explorer   |   Other: ______',
    '3. Which science concepts does your game teach? (must include ALL):',
    '   □ Planet sizes (absolute)     □ Relative scale     □ Unit conversion     □ Gravity',
    '4. Describe how the player interacts with the game:',
    '   _____________________________________________________________',
    '   _____________________________________________________________',
    '5. What makes your game fun?',
    '   _____________________________________________________________',
    '6. Write your first prompt for Claude Code:',
    '   _____________________________________________________________',
    '   _____________________________________________________________',
    '   _____________________________________________________________',
]
for q in questions:
    doc.add_paragraph(q)

doc.add_paragraph()
doc.add_heading('Appendix B: Peer Playtest Feedback', level=1)
doc.add_paragraph('Your Name: _______________  Game You Tested: _______________')
feedback_qs = [
    '1. Did the game teach you something about planets, gravity, or unit conversion?  □ Yes  □ Sort of  □ No',
    '2. What did you learn? _______________________',
    '3. What was the most fun part? _______________________',
    '4. Rate the game (circle):  ⭐  ⭐⭐  ⭐⭐⭐  ⭐⭐⭐⭐  ⭐⭐⭐⭐⭐',
    '5. One suggestion to make it better: _______________________',
]
for f in feedback_qs:
    doc.add_paragraph(f)

doc.add_paragraph()
doc.add_heading('Appendix C: VEX Go Strategy Planning Worksheet', level=1)
doc.add_paragraph('Team Name: _______________________')
strat_qs = [
    '1. Competition Task (in your own words): _______________________',
    '   _____________________________________________________________',
    '2. Our robot\'s sensors and tools:  □ Bumper  □ Eye Sensor  □ Electromagnet  □ Other: ______',
    '3. Our strategy (step by step):',
    '   Step 1: _______________________',
    '   Step 2: _______________________',
    '   Step 3: _______________________',
    '   Step 4: _______________________',
    '4. Pseudocode from Claude Code (paste or write below):',
    '   _____________________________________________________________',
    '   _____________________________________________________________',
    '   _____________________________________________________________',
    '5. After testing — what worked? _______________________',
    '6. After testing — what needs to change? _______________________',
    '7. Our #1 goal for competition: _______________________',
]
for s in strat_qs:
    doc.add_paragraph(s)

# Save
doc.save('C:/dev/sam_school/Lesson_Plan_April_15_2026.docx')
print("Document saved successfully!")
